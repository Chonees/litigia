import base64
import io

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.core.config import settings
from app.models.schemas import (
    EscritoRequest,
    EscritoResponse,
    FalloResult,
    OficioRequest,
    ResumenFalloRequest,
    ResumenFalloResponse,
)
from app.services.rag import search_jurisprudencia
from app.models.schemas import JurisprudenciaQuery

def _get_client() -> anthropic.AsyncAnthropic:
    return anthropic.AsyncAnthropic(api_key=settings.anthropic_api_key)


ESCRITO_PROMPT = """Sos un abogado argentino con 20 años de experiencia redactando escritos judiciales.
Generá un {tipo} con los siguientes datos:

DATOS DEL CASO:
- Fuero: {fuero}
- Tema: {tema}
- Posición procesal: {posicion}
- Jurisdicción: {jurisdiccion}
- Descripción: {datos_caso}

JURISPRUDENCIA DISPONIBLE (SOLO citá estos fallos, NO inventes):
{jurisprudencia}

INSTRUCCIONES:
1. Usá el formato procesal correcto para {jurisdiccion}
2. Incluí todas las secciones estándar: OBJETO, HECHOS, DERECHO, PRUEBA, PETITORIO
3. Citá SOLO la jurisprudencia proporcionada arriba — NO inventes fallos
4. Usá lenguaje jurídico formal argentino
5. Incluí artículos de ley relevantes (CPCCN, CC, LCT, etc.)
6. Terminá con "Provea V.S. de conformidad" o fórmula equivalente

Generá el escrito COMPLETO:"""


RESUMEN_PROMPT = """Sos un abogado argentino experto en análisis de jurisprudencia.
Resumí el siguiente fallo en formato estructurado.

FALLO:
{texto}

Respondé con un JSON con esta estructura exacta:
{{
  "hechos": "resumen de los hechos en 5 líneas máximo",
  "cuestion_juridica": "la cuestión jurídica en disputa",
  "argumentos_actor": "argumentos principales del actor/recurrente",
  "argumentos_demandado": "argumentos principales del demandado/recurrido",
  "resolucion": "qué resolvió el tribunal",
  "doctrina_aplicada": "doctrina jurídica aplicada",
  "articulos_citados": ["art. X del Código Y", "art. Z de la Ley N"]
}}

SOLO el JSON:"""


OFICIO_PROMPT = """Sos un abogado argentino redactando un oficio judicial.

DATOS:
- Destinatario: {destinatario}
- Motivo: {motivo}
- Expediente: {datos_expediente}
- Datos requeridos: {datos_requeridos}

Generá el oficio COMPLETO con formato correcto:
- Encabezado con lugar y fecha
- Referencia al expediente
- Solicitud formal
- Fórmula de cierre
- Espacio para firma del letrado

Oficio:"""


def _create_docx(title: str, content: str) -> str:
    """Create a .docx file and return as base64 string."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()  # spacing

    # Content - split by sections
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        para = doc.add_paragraph()

        # Detect section headers (Roman numerals, all caps, etc.)
        is_header = (
            stripped.startswith(("I.", "II.", "III.", "IV.", "V.", "VI.", "VII.", "VIII."))
            or stripped.isupper()
        )

        if is_header:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(stripped)
            run.bold = True
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return base64.b64encode(buffer.read()).decode("utf-8")


async def generate_escrito(request: EscritoRequest) -> EscritoResponse:
    """Generate a legal document with real jurisprudencia."""
    # First, search for relevant jurisprudencia
    query = JurisprudenciaQuery(
        descripcion_caso=request.datos_caso,
        jurisdiccion=request.jurisdiccion,
        fuero=request.fuero,
        top_k=5,
    )
    juris_result = await search_jurisprudencia(query)

    # Format jurisprudencia for the prompt
    juris_text = ""
    for i, fallo in enumerate(juris_result.fallos, 1):
        juris_text += (
            f"\n{i}. {fallo.caratula} ({fallo.tribunal}, {fallo.fecha})\n"
            f"   Resumen: {fallo.resumen}\n"
            f"   Cita: \"{fallo.cita_textual}\"\n"
        )

    if not juris_text:
        juris_text = "(No se encontró jurisprudencia específica. NO cites fallos inventados.)"

    response = await _get_client().messages.create(
        model=settings.anthropic_model,
        max_tokens=4000,
        messages=[
            {
                "role": "user",
                "content": ESCRITO_PROMPT.format(
                    tipo=request.tipo,
                    fuero=request.fuero,
                    tema=request.tema,
                    posicion=request.posicion,
                    jurisdiccion=request.jurisdiccion,
                    datos_caso=request.datos_caso,
                    jurisprudencia=juris_text,
                ),
            }
        ],
    )

    contenido = response.content[0].text

    # Generate .docx
    title = f"{request.tipo.upper()} - {request.tema}"
    docx_base64 = _create_docx(title, contenido)

    return EscritoResponse(
        contenido_texto=contenido,
        jurisprudencia_citada=juris_result.fallos,
        archivo_docx_base64=docx_base64,
    )


async def resumir_fallo(request: ResumenFalloRequest) -> ResumenFalloResponse:
    """Summarize a court ruling."""
    response = await _get_client().messages.create(
        model=settings.anthropic_model,
        max_tokens=2000,
        messages=[
            {
                "role": "user",
                "content": RESUMEN_PROMPT.format(texto=request.texto_fallo[:15000]),
            }
        ],
    )

    import json

    try:
        data = json.loads(response.content[0].text)
        return ResumenFalloResponse(**data)
    except (json.JSONDecodeError, KeyError):
        return ResumenFalloResponse(
            hechos="Error al procesar el fallo",
            cuestion_juridica="",
            argumentos_actor="",
            argumentos_demandado="",
            resolucion="",
            doctrina_aplicada="",
            articulos_citados=[],
        )


async def generate_oficio(request: OficioRequest) -> str:
    """Generate an oficio judicial."""
    response = await _get_client().messages.create(
        model=settings.anthropic_model,
        max_tokens=1500,
        messages=[
            {
                "role": "user",
                "content": OFICIO_PROMPT.format(
                    destinatario=request.destinatario,
                    motivo=request.motivo,
                    datos_expediente=request.datos_expediente,
                    datos_requeridos=request.datos_requeridos,
                ),
            }
        ],
    )

    return response.content[0].text
